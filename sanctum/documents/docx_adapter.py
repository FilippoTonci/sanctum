"""Word (.docx) structured reader/writer backed by python-docx.

Segment granularity is the *run* — the atomic unit that owns its own
formatting (bold, italic, color, font). Run-level round-tripping keeps
every visual property intact; the only cost is that entities which
straddle two runs (e.g. ``Alice`` bolded, trailing ``Smith`` not) won't
be detected.

Segment IDs:
    body/p{i}/r{j}                                    body paragraph run
    table/t{t}/row{r}/cell{c}/p{p}/r{j}               table cell run

Phase 1 does not read headers, footers, footnotes, or comments — they
are preserved byte-for-byte in the raw handle but not anonymized. This
is an explicit scope cut; lifting it is a Phase 2 task.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml.ns import qn

from sanctum.core.exceptions import StagedMappingParseError
from sanctum.core.models import ReviewComment, ReviewDecision
from sanctum.documents.review import format_comment_body, make_detection_id, parse_trailers
from sanctum.documents.structured import build_document, build_segment

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

    from sanctum.core.models import AnonymizationResult, StructuredDocument, TextSegment


def _iter_paragraph_runs(paragraph: Paragraph, prefix: str) -> list[tuple[str, Run]]:
    """Return ``[(segment_id, run), ...]`` for every run in ``paragraph``."""
    return [(f"{prefix}/r{j}", run) for j, run in enumerate(paragraph.runs)]


def _collect_anchor_texts(handle: DocxDocument) -> dict[int, str]:
    """Map ``comment_id`` → concatenated text anchored inside its range.

    Walks ``<w:body>`` in document order, tracking every ``commentRangeStart``
    / ``commentRangeEnd`` pair and accumulating ``<w:t>`` text that sits
    between them. Tables are traversed too because they live inside the
    body element. Ranges nest and span paragraphs, so each open range gets
    its own accumulator keyed by ``w:id``.
    """
    body = handle.element.body
    open_ranges: dict[int, list[str]] = {}
    anchor_texts: dict[int, str] = {}
    for elem in body.iter():
        tag = elem.tag
        if tag == qn("w:commentRangeStart"):
            cid = int(elem.get(qn("w:id")))
            open_ranges[cid] = []
        elif tag == qn("w:commentRangeEnd"):
            cid = int(elem.get(qn("w:id")))
            anchor_texts[cid] = "".join(open_ranges.pop(cid, []))
        elif tag == qn("w:t"):
            chunk = elem.text or ""
            for acc in open_ranges.values():
                acc.append(chunk)
    return anchor_texts


class Reader:
    """Read a .docx file into a StructuredDocument keyed by run."""

    def read(self, path: Path) -> StructuredDocument:
        doc: DocxDocument = Document(str(path))
        segments: list[TextSegment] = []

        for i, para in enumerate(doc.paragraphs):
            for seg_id, run in _iter_paragraph_runs(para, f"body/p{i}"):
                segments.append(build_segment(seg_id, run.text))

        for t, table in enumerate(doc.tables):
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    for p, para in enumerate(cell.paragraphs):
                        prefix = f"table/t{t}/row{r}/cell{c}/p{p}"
                        for seg_id, run in _iter_paragraph_runs(para, prefix):
                            segments.append(build_segment(seg_id, run.text))

        return build_document(
            source_path=path,
            fmt="docx",
            segments=segments,
            raw_handle=doc,
        )

    def read_review_decisions(self, path: Path) -> list[ReviewDecision]:
        """Classify every Word comment as accepted / rejected / user-added.

        For each comment:

        - Parse its body for ``<!-- sanctum:v=1 ... -->`` trailers.
        - If any trailer's ``replacement`` still appears in the anchored run
          text, the reviewer left Sanctum's change alone → ``accepted``.
        - If a trailer is present but its ``replacement`` is gone, the
          reviewer restored the original text → ``rejected`` (the trailer is
          kept on the decision so commit-review knows what *not* to stage).
        - A comment with no trailer is a reviewer-authored span flag →
          ``user_added``.

        Comments the reviewer deleted entirely produce no decision — and
        commit-review takes no action for unreconciled detections, which is
        the correct conservative default. Malformed trailers fall back to
        ``user_added`` rather than raising so a single broken comment can't
        block the whole commit.
        """
        handle: DocxDocument = Document(str(path))
        anchor_texts = _collect_anchor_texts(handle)
        decisions: list[ReviewDecision] = []

        for comment in sorted(handle.comments, key=lambda c: c.comment_id):
            body = comment.text
            anchor = anchor_texts.get(comment.comment_id, "")
            try:
                trailers = parse_trailers(body)
            except StagedMappingParseError:
                trailers = []
            if not trailers:
                decisions.append(
                    ReviewDecision(
                        kind="user_added",
                        user_comment_body=body,
                        user_anchor_text=anchor,
                    )
                )
                continue
            for trailer in trailers:
                kind = "accepted" if trailer.replacement in anchor else "rejected"
                decisions.append(
                    ReviewDecision(
                        kind=kind,
                        staged=trailer,
                    )
                )

        return decisions


class Writer:
    """Project a mutated StructuredDocument back to a .docx file.

    Mutates the raw python-docx handle in place by matching segment IDs
    back to their owning runs, then saves to ``path``. A missing
    ``raw_handle`` is a programmer error — the document must come from
    the matching :class:`Reader`.
    """

    def write(self, doc: StructuredDocument, path: Path) -> None:
        handle: DocxDocument | None = doc.raw_handle
        if handle is None:
            raise ValueError(
                "DocxWriter requires the StructuredDocument.raw_handle from DocxReader"
            )

        run_index = self._build_run_index(handle)
        for segment in doc.segments:
            run = run_index.get(segment.id)
            if run is None:
                continue
            run.text = segment.text

        handle.save(str(path))

    def emit_review(
        self,
        doc: StructuredDocument,
        path: Path,
        results_by_segment: dict[str, AnonymizationResult],
    ) -> None:
        """Write the document + one Word comment per detection.

        Mutates the raw handle in the same way as :meth:`write` — anonymized
        text replaces each run — then attaches a native Word comment to the
        owning run for every detection in ``results_by_segment``. Existing
        comments on the input document pass through unmodified because we
        only ever add, never walk/rewrite the input's comments collection.

        Each emitted comment body is:

            Sanctum applied 'replace' to PERSON (score 0.92):
            "John Smith" → "[PERSON_1]".
            To reject: restore the original text and delete this comment.
            <!-- sanctum:v=1 detection_id=... -->

        The trailer is what ``commit-review`` parses back via
        :meth:`Reader.read_review_decisions`. ``detection_id`` is a content
        hash of ``(entity_type, original, "<segment_id>:<i>")`` — stable
        across re-runs on the same input so copy-paste in Word doesn't
        break reconciliation.

        Raises ``ValueError`` if the engine handed us a result without
        ``per_detection_replacements`` populated — we refuse to guess the
        replacement text rather than emit a misleading trailer.
        """
        handle: DocxDocument | None = doc.raw_handle
        if handle is None:
            raise ValueError(
                "DocxWriter requires the StructuredDocument.raw_handle from DocxReader"
            )

        run_index = self._build_run_index(handle)
        for segment in doc.segments:
            run = run_index.get(segment.id)
            if run is None:
                continue
            run.text = segment.text

        for segment_id, result in results_by_segment.items():
            run = run_index.get(segment_id)
            if run is None or not result.detections:
                continue
            replacements = result.per_detection_replacements
            if replacements is None or len(replacements) != len(result.detections):
                raise ValueError(
                    f"emit_review requires per_detection_replacements aligned with "
                    f"detections for segment {segment_id!r}; got {replacements!r}"
                )
            for i, detection in enumerate(result.detections):
                comment = ReviewComment(
                    detection_id=make_detection_id(
                        detection.entity_type, detection.text_span, f"{segment_id}:{i}"
                    ),
                    entity_type=detection.entity_type,
                    score=detection.score,
                    original=detection.text_span,
                    replacement=replacements[i],
                    operator=result.operators_applied.get(detection.entity_type, ""),
                )
                handle.add_comment(
                    [run],
                    text=format_comment_body(comment),
                    author="Sanctum",
                    initials="S",
                )

        handle.save(str(path))

    @staticmethod
    def _build_run_index(handle: DocxDocument) -> dict[str, Run]:
        """Rebuild the same id → run map the reader produced."""
        index: dict[str, Run] = {}

        for i, para in enumerate(handle.paragraphs):
            for seg_id, run in _iter_paragraph_runs(para, f"body/p{i}"):
                index[seg_id] = run

        for t, table in enumerate(handle.tables):
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    for p, para in enumerate(cell.paragraphs):
                        prefix = f"table/t{t}/row{r}/cell{c}/p{p}"
                        for seg_id, run in _iter_paragraph_runs(para, prefix):
                            index[seg_id] = run

        return index
