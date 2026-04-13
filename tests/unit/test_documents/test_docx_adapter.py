"""Unit tests for the .docx reader/writer adapter pair."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from sanctum.documents.docx_adapter import Reader, Writer

FIXTURE = Path("tests/fixtures/office/nda_contract.docx")


@pytest.fixture()
def reader() -> Reader:
    return Reader()


@pytest.fixture()
def writer() -> Writer:
    return Writer()


def test_reader_produces_run_level_segments(reader: Reader) -> None:
    doc = reader.read(FIXTURE)
    assert doc.format == "docx"
    assert doc.source_path == FIXTURE
    assert doc.segments, "expected at least one run in the NDA fixture"
    assert all(seg.id.startswith(("body/", "table/")) for seg in doc.segments)


def test_reader_keeps_raw_handle(reader: Reader) -> None:
    doc = reader.read(FIXTURE)
    assert doc.raw_handle is not None


def test_round_trip_preserves_text(reader: Reader, writer: Writer, tmp_path: Path) -> None:
    """Read → write without mutation ⇒ same visible text."""
    doc = reader.read(FIXTURE)
    out = tmp_path / "round_trip.docx"
    writer.write(doc, out)

    original = Document(str(FIXTURE))
    written = Document(str(out))
    assert [p.text for p in original.paragraphs] == [p.text for p in written.paragraphs]


def test_writer_applies_segment_edits(reader: Reader, writer: Writer, tmp_path: Path) -> None:
    doc = reader.read(FIXTURE)
    target = next(seg for seg in doc.segments if seg.text.strip())
    new_segments = [
        seg.model_copy(update={"text": "REDACTED"}) if seg.id == target.id else seg
        for seg in doc.segments
    ]
    mutated = doc.model_copy(update={"segments": new_segments})
    mutated.raw_handle = doc.raw_handle

    out = tmp_path / "edited.docx"
    writer.write(mutated, out)

    reread = reader.read(out)
    assert any(seg.id == target.id and seg.text == "REDACTED" for seg in reread.segments)


def test_writer_rejects_missing_raw_handle(writer: Writer, reader: Reader, tmp_path: Path) -> None:
    doc = reader.read(FIXTURE)
    doc.raw_handle = None
    with pytest.raises(ValueError, match="raw_handle"):
        writer.write(doc, tmp_path / "x.docx")


def test_registry_dispatches_to_docx_adapter() -> None:
    from sanctum.documents import adapter_for

    r, w = adapter_for(FIXTURE)
    assert isinstance(r, Reader)
    assert isinstance(w, Writer)


def test_round_trips_table_cells(reader: Reader, writer: Writer, tmp_path: Path) -> None:
    """Tables have their own segment-id namespace and must round-trip."""
    src = tmp_path / "with_table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Alice"
    table.rows[0].cells[1].text = "Bob"
    table.rows[1].cells[0].text = "Carol"
    table.rows[1].cells[1].text = "Dan"
    doc.save(str(src))

    parsed = reader.read(src)
    table_segments = [s for s in parsed.segments if s.id.startswith("table/")]
    assert {s.text for s in table_segments} == {"Alice", "Bob", "Carol", "Dan"}

    target = next(s for s in parsed.segments if s.text == "Carol")
    new_segments = [
        s.model_copy(update={"text": "<PERSON>"}) if s.id == target.id else s
        for s in parsed.segments
    ]
    mutated = parsed.model_copy(update={"segments": new_segments})
    mutated.raw_handle = parsed.raw_handle

    out = tmp_path / "out.docx"
    writer.write(mutated, out)
    rewritten = Document(str(out))
    cell_text = {cell.text for row in rewritten.tables[0].rows for cell in row.cells}
    assert "<PERSON>" in cell_text
    assert "Carol" not in cell_text
