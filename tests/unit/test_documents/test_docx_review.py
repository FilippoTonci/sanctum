"""Unit tests for `DocxWriter.emit_review` (Phase 1.5 WS2).

These tests construct `AnonymizationResult` objects directly — they do
not run Presidio — so review emission is exercised in isolation from
the detection stack. The shape of the `AnonymizationResult` input is
the contract between the engine and the adapter; pinning it here keeps
WS6's reconciliation work decoupled from Presidio version drift.
"""

from __future__ import annotations

import re
from pathlib import Path

import pytest
from docx import Document
from sanctum.core.models import (
    AnonymizationResult,
    DetectionResult,
    StructuredDocument,
    TextSegment,
)
from sanctum.core.protocols import ReviewEmittingWriter
from sanctum.documents.docx_adapter import Reader, Writer
from sanctum.documents.review import parse_trailer
from sanctum.documents.structured import build_document


def _segment_doc(path: Path, text: str) -> tuple[StructuredDocument, TextSegment]:
    """Build a one-paragraph .docx and its StructuredDocument."""
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))
    parsed = Reader().read(path)
    segment = next(s for s in parsed.segments if s.text == text)
    return parsed, segment


def _anonymized_doc(
    parsed: StructuredDocument, segment: TextSegment, new_text: str
) -> StructuredDocument:
    """Swap one segment's text for ``new_text`` and keep the raw handle."""
    new_segments = [
        s.model_copy(update={"text": new_text}) if s.id == segment.id else s
        for s in parsed.segments
    ]
    mutated = parsed.model_copy(update={"segments": new_segments})
    mutated.raw_handle = parsed.raw_handle
    return mutated


def test_writer_satisfies_review_emitting_writer_protocol() -> None:
    """isinstance guard in the engine must accept the Writer."""
    assert isinstance(Writer(), ReviewEmittingWriter)


def test_emit_review_attaches_one_comment_per_detection(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    parsed, segment = _segment_doc(src, "Alice went to Paris")
    mutated = _anonymized_doc(parsed, segment, "<PERSON> went to <LOCATION>")

    result = AnonymizationResult(
        original_text="Alice went to Paris",
        anonymized_text="<PERSON> went to <LOCATION>",
        detections=[
            DetectionResult(entity_type="PERSON", start=0, end=5, score=0.92, text_span="Alice"),
            DetectionResult(
                entity_type="LOCATION", start=14, end=19, score=0.88, text_span="Paris"
            ),
        ],
        operators_applied={"PERSON": "replace", "LOCATION": "replace"},
        per_detection_replacements=["<PERSON>", "<LOCATION>"],
    )

    out = tmp_path / "out.docx"
    Writer().emit_review(mutated, out, {segment.id: result})

    read_back = Document(str(out))
    comments = list(read_back.comments)
    assert len(comments) == 2
    assert all(c.author == "Sanctum" for c in comments)
    # Trailer round-trips every field we care about.
    trailers = [parse_trailer(c.text) for c in comments]
    by_entity = {t.entity_type: t for t in trailers}
    assert by_entity["PERSON"].original == "Alice"
    assert by_entity["PERSON"].replacement == "<PERSON>"
    assert by_entity["PERSON"].operator == "replace"
    assert by_entity["LOCATION"].original == "Paris"
    assert by_entity["LOCATION"].replacement == "<LOCATION>"


def test_emit_review_anonymized_text_lands_in_run(tmp_path: Path) -> None:
    """The run must carry the anonymized text, not the original — a
    reviewer reading the file sees '<PERSON>' not 'Alice'."""
    src = tmp_path / "in.docx"
    parsed, segment = _segment_doc(src, "Alice works here")
    mutated = _anonymized_doc(parsed, segment, "<PERSON> works here")

    result = AnonymizationResult(
        original_text="Alice works here",
        anonymized_text="<PERSON> works here",
        detections=[
            DetectionResult(entity_type="PERSON", start=0, end=5, score=0.9, text_span="Alice")
        ],
        operators_applied={"PERSON": "replace"},
        per_detection_replacements=["<PERSON>"],
    )

    out = tmp_path / "out.docx"
    Writer().emit_review(mutated, out, {segment.id: result})

    read_back = Document(str(out))
    body_text = "\n".join(p.text for p in read_back.paragraphs)
    assert "Alice" not in body_text
    assert "<PERSON>" in body_text


def test_emit_review_with_no_detections_writes_no_comments(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    parsed, _segment = _segment_doc(src, "Clean text, no PII")

    out = tmp_path / "out.docx"
    Writer().emit_review(parsed, out, {})

    read_back = Document(str(out))
    assert len(list(read_back.comments)) == 0


def test_emit_review_preserves_existing_comments(tmp_path: Path) -> None:
    """The fixture has a pre-existing reviewer comment; emit_review must
    pass it through. A stomp would silently nuke real review history."""
    src = tmp_path / "with_existing.docx"
    doc = Document()
    para = doc.add_paragraph("")
    run = para.add_run("Alice is here")
    doc.add_comment([run], text="reviewer note", author="Jane", initials="J")
    doc.save(str(src))
    parsed = Reader().read(src)

    segment = next(s for s in parsed.segments if s.text == "Alice is here")
    mutated = _anonymized_doc(parsed, segment, "<PERSON> is here")

    result = AnonymizationResult(
        original_text="Alice is here",
        anonymized_text="<PERSON> is here",
        detections=[
            DetectionResult(entity_type="PERSON", start=0, end=5, score=0.9, text_span="Alice")
        ],
        operators_applied={"PERSON": "replace"},
        per_detection_replacements=["<PERSON>"],
    )

    out = tmp_path / "out.docx"
    Writer().emit_review(mutated, out, {segment.id: result})

    read_back = Document(str(out))
    authors = {c.author for c in read_back.comments}
    assert {"Jane", "Sanctum"} <= authors
    # Existing comment text is intact.
    assert any("reviewer note" in c.text for c in read_back.comments)


def test_emit_review_refuses_without_per_detection_replacements(tmp_path: Path) -> None:
    """An upstream adapter that doesn't surface per-detection replacements
    must not silently produce misleading trailers."""
    src = tmp_path / "in.docx"
    parsed, segment = _segment_doc(src, "Alice")
    mutated = _anonymized_doc(parsed, segment, "<PERSON>")

    result = AnonymizationResult(
        original_text="Alice",
        anonymized_text="<PERSON>",
        detections=[
            DetectionResult(entity_type="PERSON", start=0, end=5, score=0.9, text_span="Alice")
        ],
        operators_applied={"PERSON": "replace"},
        per_detection_replacements=None,
    )

    with pytest.raises(ValueError, match="per_detection_replacements"):
        Writer().emit_review(mutated, tmp_path / "out.docx", {segment.id: result})


def test_emit_review_requires_raw_handle() -> None:
    """Mirror of the write() contract: handle is mandatory."""
    doc = build_document(
        source_path=Path("/nonexistent.docx"),
        fmt="docx",
        segments=[],
        raw_handle=None,
    )
    with pytest.raises(ValueError, match="raw_handle"):
        Writer().emit_review(doc, Path("/tmp/x.docx"), {})


def test_detection_id_is_stable_for_same_input(tmp_path: Path) -> None:
    """Content-addressed ids must match across emissions on identical
    input — the reconciliation in WS6 relies on it surviving copy-paste."""
    src = tmp_path / "in.docx"
    parsed, segment = _segment_doc(src, "Alice")
    mutated = _anonymized_doc(parsed, segment, "<PERSON>")

    result = AnonymizationResult(
        original_text="Alice",
        anonymized_text="<PERSON>",
        detections=[
            DetectionResult(entity_type="PERSON", start=0, end=5, score=0.9, text_span="Alice")
        ],
        operators_applied={"PERSON": "replace"},
        per_detection_replacements=["<PERSON>"],
    )

    def _extract_id(path: Path) -> str:
        (comment,) = list(Document(str(path)).comments)
        match = re.search(r"detection_id=([0-9a-f]+)", comment.text)
        assert match is not None
        return match.group(1)

    out1 = tmp_path / "o1.docx"
    Writer().emit_review(mutated, out1, {segment.id: result})
    # Re-read a fresh handle so we emit from a clean state.
    parsed2, segment2 = _segment_doc(tmp_path / "in2.docx", "Alice")
    mutated2 = _anonymized_doc(parsed2, segment2, "<PERSON>")
    out2 = tmp_path / "o2.docx"
    Writer().emit_review(mutated2, out2, {segment2.id: result})

    assert _extract_id(out1) == _extract_id(out2)
