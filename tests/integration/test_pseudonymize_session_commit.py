"""Pseudonymize end-to-end through a review session (Phase 1.5 WS4).

Boots a real waitress server with real Presidio engines, runs full
review-session flows against the DOCX NDA fixture under the
``pseudonymize`` operator, and asserts on what actually reaches the
encrypted mapping store. Covers the WS4 plan's six verification
points:

1. Creating a session does not mint into the store.
2. GET returns previews with plausible pseudonyms.
3. PATCHing mixed decisions (accept-default, reject, per-entity
   override, user-added) leaves non-persisting previews visible.
4. Commit persists exactly the accepted + user-added pseudonymize
   decisions and writes a trailer-free file.
5. A second session on the same entity reuses the first session's
   committed pseudonym.
6. Abandoning a session leaves the store untouched.

Each test owns its own encrypted store (function-scoped fixture) so
tests stay independent and the API's flock never collides with the
on-disk assertion path.

Marked ``integration`` so ``pytest -m "not integration"`` skips it.
"""

from __future__ import annotations

import json
import socket
import threading
import time
import urllib.error
import urllib.request
from collections.abc import Iterator
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from sanctum.analyzer.adapter import PresidioAnalyzer
from sanctum.anonymizer.adapter import PresidioAnonymizer
from sanctum.api.app import create_app
from sanctum.core.engine import SanctumEngine
from sanctum.core.review.store import SessionStore
from sanctum.security.mapping_store import EncryptedFileMappingStore
from waitress.server import create_server  # type: ignore[import-untyped]

pytestmark = pytest.mark.integration

_TOKEN = "integration-token-do-not-reuse"
_PASSPHRASE = "passw0rd-ws4"
_FIXTURE = Path(__file__).resolve().parents[1] / "fixtures" / "office" / "nda_contract.docx"


def _free_port() -> int:
    with socket.socket() as s:
        s.bind(("127.0.0.1", 0))
        return int(s.getsockname()[1])


def _host_header(base_url: str) -> str:
    return base_url.removeprefix("http://").removeprefix("https://")


def _wait_for_health(base_url: str, timeout: float = 30.0) -> None:
    deadline = time.time() + timeout
    last_exc: Exception | None = None
    while time.time() < deadline:
        try:
            req = urllib.request.Request(
                f"{base_url}/health",
                headers={"Host": _host_header(base_url)},
            )
            with urllib.request.urlopen(req, timeout=1.0) as r:
                if r.status == 200:
                    return
        except (urllib.error.URLError, ConnectionError) as exc:
            last_exc = exc
        time.sleep(0.05)
    raise RuntimeError(f"server never came up: {last_exc!r}")


def _cheap_factory(path: Path) -> EncryptedFileMappingStore:
    """KDF tuned for test speed — same params the unit suite uses."""
    return EncryptedFileMappingStore(
        path,
        kdf_time_cost=1,
        kdf_memory_cost=8,
        kdf_parallelism=1,
    )


def _request(
    method: str,
    url: str,
    *,
    token: str | None = None,
    body: dict[str, Any] | None = None,
) -> tuple[int, Any]:
    netloc = url.split("/", 3)[2]
    headers: dict[str, str] = {"Host": netloc}
    if token is not None:
        headers["Authorization"] = f"Bearer {token}"
    data: bytes | None = None
    if body is not None:
        headers["Content-Type"] = "application/json"
        data = json.dumps(body).encode()
    req = urllib.request.Request(url, data=data, method=method, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=15.0) as r:
            raw = r.read().decode()
            payload = json.loads(raw) if raw else ""
            return r.status, payload
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode()
        payload = json.loads(raw) if raw else {}
        return exc.code, payload


@pytest.fixture(scope="module")
def server(tmp_path_factory: pytest.TempPathFactory) -> Iterator[tuple[str, str]]:
    """Waitress + real Presidio + tmp-rooted SessionStore + cheap-KDF mapping factory.

    Scoped module-wide because Presidio engine boot is slow (~3s each).
    Individual encrypted stores are per-test (function-scoped fixture
    below) so the API's flock never collides with on-disk assertions.
    """
    if not _FIXTURE.is_file():
        pytest.skip(f"fixture missing: {_FIXTURE}")

    port = _free_port()
    sessions_root = tmp_path_factory.mktemp("sessions")
    engine = SanctumEngine(
        analyzer=PresidioAnalyzer(),
        anonymizer=PresidioAnonymizer(),
    )
    app = create_app(
        token=_TOKEN,
        host="127.0.0.1",
        port=port,
        engine=engine,
        mapping_store_factory=_cheap_factory,
        session_store=SessionStore(root=sessions_root),
    )
    wsgi_server = create_server(app, host="127.0.0.1", port=port, threads=2)
    thread = threading.Thread(target=wsgi_server.run, name="sanctum-pseudo-waitress", daemon=True)
    thread.start()

    base_url = f"http://127.0.0.1:{port}"
    try:
        _wait_for_health(base_url)
        yield base_url, _TOKEN
    finally:
        wsgi_server.close()
        thread.join(timeout=5.0)


@pytest.fixture
def store_path(tmp_path: Path, server: tuple[str, str]) -> Iterator[Path]:
    """Fresh encrypted store per test; auto-locks on teardown so the
    next test (or the on-disk assertion) can grab the flock.
    """
    base, token = server
    path = tmp_path / "store.bin"
    yield path
    # Best-effort: if the test left the store unlocked, release the
    # flock now so test-scoped on-disk reads don't hang.
    _request("POST", f"{base}/mapping/lock", token=token)


def _unlock(base: str, token: str, path: Path) -> None:
    status, body = _request(
        "POST",
        f"{base}/mapping/unlock",
        token=token,
        body={"store_path": str(path), "passphrase": _PASSPHRASE},
    )
    assert status == 200, body


def _lock(base: str, token: str) -> None:
    status, _ = _request("POST", f"{base}/mapping/lock", token=token)
    assert status == 200


def _docx_full_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def _read_store_entries(path: Path) -> dict[str, dict[str, str]]:
    """Open the encrypted store directly and return a copy of its entries.

    Call this *only* after the API's ``/mapping/lock`` has released the
    flock, otherwise the ``_acquire_flock`` inside ``unlock`` will
    raise ``MappingStoreError``.
    """
    if not path.exists():
        return {}
    store = _cheap_factory(path)
    store.unlock(_PASSPHRASE)
    try:
        # Private attr: the test needs raw state so it can prove
        # "nothing was written" for originals it didn't know about.
        entries = store._entries or {}
        return {et: dict(bucket) for et, bucket in entries.items()}
    finally:
        store.lock()


# ---------- tests ----------


def test_creating_session_does_not_mint_into_store(
    server: tuple[str, str], store_path: Path
) -> None:
    """Session create seeds previews without writing to the store."""
    base, token = server
    _unlock(base, token, store_path)
    try:
        status, body = _request(
            "POST",
            f"{base}/review-sessions",
            token=token,
            body={
                "input_path": str(_FIXTURE),
                "default_operator": "pseudonymize",
            },
        )
        assert status == 201, body
        assert body["default_operator"] == "pseudonymize"
        assert len(body["proposals"]) > 0

        for proposal in body["proposals"]:
            preview = body["previews"][proposal["detection_id"]]
            assert preview
            if proposal["entity_type"] in {"PERSON", "LOCATION", "ORGANIZATION"}:
                assert preview != proposal["original"]
    finally:
        _lock(base, token)

    # Nothing should have been persisted.
    assert _read_store_entries(store_path) == {}


def test_commit_persists_only_accepted_and_user_added_pseudonyms(
    server: tuple[str, str], store_path: Path, tmp_path: Path
) -> None:
    """Mixed decisions → commit populates the store with exactly the
    accepted + user-added pseudonymize entries. Final file trailer-free.
    """
    base, token = server
    _unlock(base, token, store_path)
    accepted: dict[str, Any]
    rejected: dict[str, Any]
    overridden: dict[str, Any]
    ua_original: str
    try:
        _, created = _request(
            "POST",
            f"{base}/review-sessions",
            token=token,
            body={
                "input_path": str(_FIXTURE),
                "default_operator": "pseudonymize",
            },
        )
        session_id = created["id"]
        proposals = created["proposals"]
        persons = [p for p in proposals if p["entity_type"] == "PERSON"]
        assert len(persons) >= 3, f"need ≥3 PERSON proposals, got {len(persons)}"
        accepted = persons[0]
        rejected = persons[1]
        overridden = persons[2]

        _, ab = _request(
            "PATCH",
            f"{base}/review-sessions/{session_id}/decisions/{accepted['detection_id']}",
            token=token,
            body={"status": "accept"},
        )
        assert ab["decision"]["status"] == "accept"

        _request(
            "PATCH",
            f"{base}/review-sessions/{session_id}/decisions/{rejected['detection_id']}",
            token=token,
            body={"status": "reject"},
        )

        _, ov = _request(
            "PATCH",
            f"{base}/review-sessions/{session_id}/decisions/{overridden['detection_id']}",
            token=token,
            body={"status": "accept", "operator": "replace"},
        )
        assert ov["preview"] == "<PERSON>"

        anchor = next(s for s in created["segments"] if s["text"].strip())
        ua_original = anchor["text"].split()[0]
        status, ua_body = _request(
            "POST",
            f"{base}/review-sessions/{session_id}/decisions/user-added",
            token=token,
            body={
                "segment_anchor": anchor["id"],
                "entity_type": "PERSON",
                "original": ua_original,
            },
        )
        assert status == 201, ua_body

        # Reject every other proposal so the commit set is exactly
        # {accepted (pseudonymize), overridden (replace), user-added}.
        for p in proposals:
            pid = p["detection_id"]
            if pid in (
                accepted["detection_id"],
                rejected["detection_id"],
                overridden["detection_id"],
            ):
                continue
            _request(
                "PATCH",
                f"{base}/review-sessions/{session_id}/decisions/{pid}",
                token=token,
                body={"status": "reject"},
            )

        out_path = tmp_path / "committed.docx"
        status, body = _request(
            "POST",
            f"{base}/review-sessions/{session_id}/commit",
            token=token,
            body={"output_path": str(out_path), "attested": True},
        )
        assert status == 200, body
        assert out_path.is_file()

        text = _docx_full_text(out_path)
        assert "sanctum:" not in text
        assert "<PERSON>" in text  # from the replace override
        assert rejected["original"] in text
    finally:
        _lock(base, token)

    # Exactly the accepted-pseudonymize and user-added originals land
    # in the store. Rejected + replace-overridden must not appear.
    on_disk = _read_store_entries(store_path)
    person_bucket = on_disk.get("PERSON", {})
    assert accepted["original"] in person_bucket
    assert ua_original in person_bucket
    assert rejected["original"] not in person_bucket
    assert overridden["original"] not in person_bucket


def test_second_session_reuses_existing_pseudonym(
    server: tuple[str, str], store_path: Path, tmp_path: Path
) -> None:
    """Re-running against an already-committed original yields the same pseudonym."""
    base, token = server

    # --- first session: commit one PERSON to populate the store ---
    _unlock(base, token, store_path)
    first_original: str
    first_pseudonym: str
    try:
        _, created = _request(
            "POST",
            f"{base}/review-sessions",
            token=token,
            body={
                "input_path": str(_FIXTURE),
                "default_operator": "pseudonymize",
            },
        )
        session_id = created["id"]
        persons = [p for p in created["proposals"] if p["entity_type"] == "PERSON"]
        first_pick = persons[0]
        first_original = first_pick["original"]
        first_pseudonym = created["previews"][first_pick["detection_id"]]

        _request(
            "PATCH",
            f"{base}/review-sessions/{session_id}/decisions/{first_pick['detection_id']}",
            token=token,
            body={"status": "accept"},
        )
        for p in created["proposals"]:
            if p["detection_id"] == first_pick["detection_id"]:
                continue
            _request(
                "PATCH",
                f"{base}/review-sessions/{session_id}/decisions/{p['detection_id']}",
                token=token,
                body={"status": "reject"},
            )
        status, _ = _request(
            "POST",
            f"{base}/review-sessions/{session_id}/commit",
            token=token,
            body={"output_path": str(tmp_path / "first.docx"), "attested": True},
        )
        assert status == 200
    finally:
        _lock(base, token)

    # Sanity — the first commit actually persisted.
    on_disk = _read_store_entries(store_path)
    assert on_disk.get("PERSON", {}).get(first_original) == first_pseudonym

    # --- second session: the same proposal's preview must reuse the committed value ---
    _unlock(base, token, store_path)
    try:
        _, second = _request(
            "POST",
            f"{base}/review-sessions",
            token=token,
            body={
                "input_path": str(_FIXTURE),
                "default_operator": "pseudonymize",
            },
        )
        match = next(
            (p for p in second["proposals"] if p["original"] == first_original),
            None,
        )
        assert match is not None
        assert second["previews"][match["detection_id"]] == first_pseudonym
        _request("DELETE", f"{base}/review-sessions/{second['id']}", token=token)
    finally:
        _lock(base, token)


def test_abandon_leaves_store_untouched(server: tuple[str, str], store_path: Path) -> None:
    """Abandon deletes the session dir but does not write to the mapping store."""
    base, token = server
    _unlock(base, token, store_path)
    try:
        _, created = _request(
            "POST",
            f"{base}/review-sessions",
            token=token,
            body={
                "input_path": str(_FIXTURE),
                "default_operator": "pseudonymize",
            },
        )
        session_id = created["id"]
        for p in created["proposals"][:3]:
            _request(
                "PATCH",
                f"{base}/review-sessions/{session_id}/decisions/{p['detection_id']}",
                token=token,
                body={"status": "accept"},
            )
        status, _ = _request("DELETE", f"{base}/review-sessions/{session_id}", token=token)
        assert status == 204
    finally:
        _lock(base, token)

    # Abandon must not write; the store file may exist from the unlock
    # cycle but its entries dict should still be empty.
    assert _read_store_entries(store_path) == {}


def test_double_commit_404s(server: tuple[str, str], store_path: Path, tmp_path: Path) -> None:
    """A second commit on the same session id is refused.

    Session dir is torn down on successful commit, so the second call
    404s. Same "one commit" invariant as the WS2 unit suite — just
    enforced via the store's existence check rather than the status
    field.
    """
    base, token = server
    _unlock(base, token, store_path)
    try:
        _, created = _request(
            "POST",
            f"{base}/review-sessions",
            token=token,
            body={"input_path": str(_FIXTURE), "default_operator": "replace"},
        )
        session_id = created["id"]
        for p in created["proposals"]:
            _request(
                "PATCH",
                f"{base}/review-sessions/{session_id}/decisions/{p['detection_id']}",
                token=token,
                body={"status": "accept"},
            )
        out_path = tmp_path / "first.docx"
        status, _ = _request(
            "POST",
            f"{base}/review-sessions/{session_id}/commit",
            token=token,
            body={"output_path": str(out_path), "attested": True},
        )
        assert status == 200

        status, body = _request(
            "POST",
            f"{base}/review-sessions/{session_id}/commit",
            token=token,
            body={"output_path": str(tmp_path / "second.docx"), "attested": True},
        )
        assert status == 404, body
    finally:
        _lock(base, token)


def test_commit_rejects_pseudonymize_session_when_store_unavailable(
    server: tuple[str, str], store_path: Path, tmp_path: Path
) -> None:
    """Commit-time pseudonymize without an unlocked store → 400, not a
    cryptic downstream error.

    After ``/mapping/lock`` the API's ``SANCTUM_MAPPING_STORE`` slot is
    cleared, matching the "never unlocked this session" state; the
    route's eager check surfaces "unlock first" rather than a torch-
    bearing stack trace out of the anonymizer.
    """
    base, token = server
    _unlock(base, token, store_path)
    try:
        _, created = _request(
            "POST",
            f"{base}/review-sessions",
            token=token,
            body={
                "input_path": str(_FIXTURE),
                "default_operator": "pseudonymize",
            },
        )
        session_id = created["id"]
        _request(
            "PATCH",
            f"{base}/review-sessions/{session_id}/decisions/"
            f"{created['proposals'][0]['detection_id']}",
            token=token,
            body={"status": "accept"},
        )
    finally:
        _lock(base, token)

    status, body = _request(
        "POST",
        f"{base}/review-sessions/{session_id}/commit",
        token=token,
        body={"output_path": str(tmp_path / "locked.docx"), "attested": True},
    )
    assert status == 400, body
    assert "unlock" in body["error"].lower()

    # Clean up so the session dir does not leak into the next test.
    _unlock(base, token, store_path)
    _request("DELETE", f"{base}/review-sessions/{session_id}", token=token)
    _lock(base, token)
