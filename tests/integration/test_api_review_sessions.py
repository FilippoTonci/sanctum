"""End-to-end integration tests for the /review-sessions API.

Boots a real waitress server with real Presidio engines, runs a full
review-session flow against the DOCX NDA fixture, and asserts on the
final committed document.

The goal is to catch wiring mistakes the unit suite can't: the real
SessionStore writing to disk under 0700/0600 perms, Presidio's analyzer
actually producing the detections the API expects, the tempfile-round-
trip inside ``commit_review_session`` preserving docx formatting, and
the final file being leakage-free (``grep sanctum:`` → zero hits, the
WS2 plan's explicit invariant).

Marked ``integration`` so ``pytest -m "not integration"`` skips it.
"""

from __future__ import annotations

import json
import socket
import threading
import time
import urllib.error
import urllib.request
from collections.abc import Iterator
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from sanctum.analyzer.adapter import PresidioAnalyzer
from sanctum.anonymizer.adapter import PresidioAnonymizer
from sanctum.api.app import create_app
from sanctum.core.engine import SanctumEngine
from sanctum.core.review.store import SessionStore
from waitress.server import create_server  # type: ignore[import-untyped]

pytestmark = pytest.mark.integration

_TOKEN = "integration-token-do-not-reuse"
_FIXTURE = Path(__file__).resolve().parents[1] / "fixtures" / "office" / "nda_contract.docx"


def _free_port() -> int:
    with socket.socket() as s:
        s.bind(("127.0.0.1", 0))
        return int(s.getsockname()[1])


def _host_header(base_url: str) -> str:
    return base_url.removeprefix("http://").removeprefix("https://")


def _wait_for_health(base_url: str, timeout: float = 30.0) -> None:
    deadline = time.time() + timeout
    last_exc: Exception | None = None
    while time.time() < deadline:
        try:
            req = urllib.request.Request(
                f"{base_url}/health",
                headers={"Host": _host_header(base_url)},
            )
            with urllib.request.urlopen(req, timeout=1.0) as r:
                if r.status == 200:
                    return
        except (urllib.error.URLError, ConnectionError) as exc:
            last_exc = exc
        time.sleep(0.05)
    raise RuntimeError(f"server never came up: {last_exc!r}")


def _request(
    method: str,
    url: str,
    *,
    token: str | None = None,
    body: dict[str, Any] | None = None,
) -> tuple[int, Any]:
    """Tiny urllib client. Returns ``(status, parsed_body_or_empty)``.

    Handles 204 (empty body) by returning ``""`` as the parsed value so
    callers can assert on status without unwrapping JSON that isn't there.
    """
    netloc = url.split("/", 3)[2]
    headers: dict[str, str] = {"Host": netloc}
    if token is not None:
        headers["Authorization"] = f"Bearer {token}"
    data: bytes | None = None
    if body is not None:
        headers["Content-Type"] = "application/json"
        data = json.dumps(body).encode()
    req = urllib.request.Request(url, data=data, method=method, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=15.0) as r:
            raw = r.read().decode()
            payload = json.loads(raw) if raw else ""
            return r.status, payload
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode()
        payload = json.loads(raw) if raw else {}
        return exc.code, payload


@pytest.fixture(scope="module")
def server(tmp_path_factory: pytest.TempPathFactory) -> Iterator[tuple[str, str]]:
    """Waitress + real Presidio + tmp-rooted SessionStore."""
    if not _FIXTURE.is_file():
        pytest.skip(f"fixture missing: {_FIXTURE}")

    port = _free_port()
    sessions_root = tmp_path_factory.mktemp("sessions")
    engine = SanctumEngine(
        analyzer=PresidioAnalyzer(),
        anonymizer=PresidioAnonymizer(),
    )
    app = create_app(
        token=_TOKEN,
        host="127.0.0.1",
        port=port,
        engine=engine,
        session_store=SessionStore(root=sessions_root),
    )
    wsgi_server = create_server(app, host="127.0.0.1", port=port, threads=2)
    thread = threading.Thread(target=wsgi_server.run, name="sanctum-review-waitress", daemon=True)
    thread.start()

    base_url = f"http://127.0.0.1:{port}"
    try:
        _wait_for_health(base_url)
        yield base_url, _TOKEN
    finally:
        wsgi_server.close()
        thread.join(timeout=5.0)


def _docx_full_text(path: Path) -> str:
    """Flatten every paragraph + table-cell text for substring assertions.

    The NDA fixture is paragraph-only (no tables) but we include the
    table walk anyway so the helper is usable on future fixtures.
    """
    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


# ---------- tests ----------


def test_create_session_returns_proposals_and_previews(
    server: tuple[str, str], tmp_path: Path
) -> None:
    base, token = server
    status, body = _request(
        "POST",
        f"{base}/review-sessions",
        token=token,
        body={"input_path": str(_FIXTURE), "default_operator": "replace"},
    )
    assert status == 201, body
    assert body["status"] == "open"
    assert body["default_operator"] == "replace"
    assert body["format"] == "docx"
    assert len(body["proposals"]) > 0, "NDA fixture should yield some detections"
    # Every proposal id appears as a preview key; under "replace" every
    # preview is an ``<ENTITY>`` tag.
    for proposal in body["proposals"]:
        assert proposal["detection_id"] in body["previews"]
        assert body["previews"][proposal["detection_id"]] == f"<{proposal['entity_type']}>"


def test_get_round_trips_session(server: tuple[str, str]) -> None:
    base, token = server
    _, created = _request(
        "POST",
        f"{base}/review-sessions",
        token=token,
        body={"input_path": str(_FIXTURE), "default_operator": "replace"},
    )
    session_id = created["id"]
    status, fetched = _request("GET", f"{base}/review-sessions/{session_id}", token=token)
    assert status == 200
    assert fetched["id"] == session_id
    assert fetched["previews"] == created["previews"]


def test_patch_decisions_update_previews(server: tuple[str, str]) -> None:
    """Accept with default / custom_replacement / reject — previews change."""
    base, token = server
    _, created = _request(
        "POST",
        f"{base}/review-sessions",
        token=token,
        body={"input_path": str(_FIXTURE), "default_operator": "replace"},
    )
    session_id = created["id"]
    proposals = created["proposals"]
    assert (
        len(proposals) >= 3
    ), f"NDA fixture produced too few proposals ({len(proposals)}) to test variants"

    # Accept first with default.
    pid_a = proposals[0]["detection_id"]
    status, body_a = _request(
        "PATCH",
        f"{base}/review-sessions/{session_id}/decisions/{pid_a}",
        token=token,
        body={"status": "accept"},
    )
    assert status == 200, body_a
    assert body_a["preview"] == f"<{proposals[0]['entity_type']}>"

    # Reject second — preview becomes the original text.
    pid_b = proposals[1]["detection_id"]
    status, body_b = _request(
        "PATCH",
        f"{base}/review-sessions/{session_id}/decisions/{pid_b}",
        token=token,
        body={"status": "reject"},
    )
    assert status == 200, body_b
    assert body_b["preview"] == proposals[1]["original"]

    # Custom replacement on third — literal lands in preview.
    pid_c = proposals[2]["detection_id"]
    status, body_c = _request(
        "PATCH",
        f"{base}/review-sessions/{session_id}/decisions/{pid_c}",
        token=token,
        body={"status": "accept", "custom_replacement": "[CLIENT]"},
    )
    assert status == 200, body_c
    assert body_c["preview"] == "[CLIENT]"

    # GET reflects the combined state.
    _, fetched = _request("GET", f"{base}/review-sessions/{session_id}", token=token)
    assert fetched["previews"][pid_a] == f"<{proposals[0]['entity_type']}>"
    assert fetched["previews"][pid_b] == proposals[1]["original"]
    assert fetched["previews"][pid_c] == "[CLIENT]"


def test_commit_writes_final_file_with_no_sanctum_trailers(
    server: tuple[str, str], tmp_path: Path
) -> None:
    base, token = server
    _, created = _request(
        "POST",
        f"{base}/review-sessions",
        token=token,
        body={"input_path": str(_FIXTURE), "default_operator": "replace"},
    )
    session_id = created["id"]

    # Mix decisions: accept the first PERSON with a custom literal; reject
    # the second detection; accept every other proposal with the session
    # default. Proves custom + reject + default all survive commit.
    custom_pid: str | None = None
    rejected_pid: str | None = None
    rejected_original: str | None = None
    for proposal in created["proposals"]:
        if proposal["entity_type"] == "PERSON" and custom_pid is None:
            custom_pid = proposal["detection_id"]
            continue
        if rejected_pid is None:
            rejected_pid = proposal["detection_id"]
            rejected_original = proposal["original"]

    assert (
        custom_pid is not None and rejected_pid is not None
    ), "NDA fixture should contain at least one PERSON and one other detection"

    _request(
        "PATCH",
        f"{base}/review-sessions/{session_id}/decisions/{custom_pid}",
        token=token,
        body={"status": "accept", "custom_replacement": "[CLIENT]"},
    )
    _request(
        "PATCH",
        f"{base}/review-sessions/{session_id}/decisions/{rejected_pid}",
        token=token,
        body={"status": "reject"},
    )
    for proposal in created["proposals"]:
        pid = proposal["detection_id"]
        if pid in (custom_pid, rejected_pid):
            continue
        _request(
            "PATCH",
            f"{base}/review-sessions/{session_id}/decisions/{pid}",
            token=token,
            body={"status": "accept"},
        )

    out_path = tmp_path / "committed.docx"
    status, body = _request(
        "POST",
        f"{base}/review-sessions/{session_id}/commit",
        token=token,
        body={"output_path": str(out_path), "attested": True},
    )
    assert status == 200, body
    assert body["output_path"] == str(out_path)
    assert out_path.is_file(), "commit should produce a file at output_path"

    text = _docx_full_text(out_path)
    assert "[CLIENT]" in text, "custom replacement should land verbatim"
    assert rejected_original is not None
    assert rejected_original in text, "rejected original text should survive"
    # Leakage invariant — the final file must carry no review-trailer
    # markers.  Integration tests enforce this because the unit suite
    # can't exercise the real docx writer.
    assert "sanctum:" not in text, "final file must be trailer-free"

    # Manifest survives so the desktop's Recent Sessions list keeps the
    # audit trail; status flips to ``committed`` and the input bytes are
    # shed (covered by the unit-store tests).
    status, body = _request("GET", f"{base}/review-sessions/{session_id}", token=token)
    assert status == 200
    assert body["status"] == "committed"
    assert body["committed_at"] is not None


def test_commit_requires_attested(server: tuple[str, str], tmp_path: Path) -> None:
    base, token = server
    _, created = _request(
        "POST",
        f"{base}/review-sessions",
        token=token,
        body={"input_path": str(_FIXTURE), "default_operator": "replace"},
    )
    status, body = _request(
        "POST",
        f"{base}/review-sessions/{created['id']}/commit",
        token=token,
        body={"output_path": str(tmp_path / "out.docx")},
    )
    assert status == 400
    assert "attested" in body["error"]


def test_abandon_marks_session_terminal_and_keeps_manifest(server: tuple[str, str]) -> None:
    base, token = server
    _, created = _request(
        "POST",
        f"{base}/review-sessions",
        token=token,
        body={"input_path": str(_FIXTURE), "default_operator": "replace"},
    )
    session_id = created["id"]
    status, _ = _request("DELETE", f"{base}/review-sessions/{session_id}", token=token)
    assert status == 204
    # Manifest survives — the desktop's Recent Sessions list reads it
    # to show what the user decided to skip.
    status, body = _request("GET", f"{base}/review-sessions/{session_id}", token=token)
    assert status == 200
    assert body["status"] == "abandoned"


def test_listing_includes_committed_and_abandoned_sessions(
    server: tuple[str, str], tmp_path: Path
) -> None:
    """Regression test: terminal sessions must keep appearing in the
    listing endpoint that powers the desktop's Recent Sessions UI.

    Pre-fix, both abandon and commit deleted the on-disk manifest, so
    the listing collapsed to "open sessions only" — which silently
    erased the user's audit trail every time they closed a session.
    """
    base, token = server
    # One session that we'll abandon.
    _, abandoned = _request(
        "POST",
        f"{base}/review-sessions",
        token=token,
        body={"input_path": str(_FIXTURE), "default_operator": "replace"},
    )
    _request("DELETE", f"{base}/review-sessions/{abandoned['id']}", token=token)

    # One session that we'll commit (after accepting all proposals).
    _, committed = _request(
        "POST",
        f"{base}/review-sessions",
        token=token,
        body={"input_path": str(_FIXTURE), "default_operator": "replace"},
    )
    for proposal in committed["proposals"]:
        _request(
            "PATCH",
            f"{base}/review-sessions/{committed['id']}/decisions/{proposal['detection_id']}",
            token=token,
            body={"status": "accept"},
        )
    _request(
        "POST",
        f"{base}/review-sessions/{committed['id']}/commit",
        token=token,
        body={"output_path": str(tmp_path / "out.docx"), "attested": True},
    )

    status, body = _request("GET", f"{base}/review-sessions", token=token)
    assert status == 200
    by_id = {s["id"]: s for s in body["sessions"]}
    assert by_id[abandoned["id"]]["status"] == "abandoned"
    assert by_id[committed["id"]]["status"] == "committed"


def test_process_file_review_true_creates_session_and_is_retrievable(
    server: tuple[str, str],
) -> None:
    """/process-file with review=true is the CLI/API shortcut: one POST to
    stage a session, then the standard /review-sessions flow takes over.

    Proves (a) the handoff shape — {session_id, review_url} — is what the
    WS3 UI will hang off of, and (b) the session id the route returned
    actually resolves via GET /review-sessions/{id}. Without (b) the
    shortcut would be a dead pointer.
    """
    base, token = server
    status, body = _request(
        "POST",
        f"{base}/process-file",
        token=token,
        body={"input_path": str(_FIXTURE), "review": True, "operator": "replace"},
    )
    assert status == 201, body
    session_id = body["session_id"]
    assert session_id
    assert body["review_url"].endswith(f"/review/{session_id}")

    status, session = _request("GET", f"{base}/review-sessions/{session_id}", token=token)
    assert status == 200, session
    assert session["id"] == session_id
    assert session["default_operator"] == "replace"
    assert len(session["proposals"]) > 0


def test_user_added_decision_round_trip(server: tuple[str, str]) -> None:
    """Add a user-added span against a real segment, then delete it."""
    base, token = server
    _, created = _request(
        "POST",
        f"{base}/review-sessions",
        token=token,
        body={"input_path": str(_FIXTURE), "default_operator": "replace"},
    )
    session_id = created["id"]
    # Any non-empty segment will do — we just need a valid anchor and
    # any substring of that segment's text for ``original``.
    anchor_segment = next(s for s in created["segments"] if s["text"].strip())
    original = anchor_segment["text"].split()[0]  # first word — guaranteed substring.
    start = anchor_segment["text"].index(original)

    status, body = _request(
        "POST",
        f"{base}/review-sessions/{session_id}/decisions/user-added",
        token=token,
        body={
            "segment_anchor": anchor_segment["id"],
            "entity_type": "PERSON",
            "original": original,
            "start": start,
            "end": start + len(original),
        },
    )
    assert status == 201, body
    ua_id = body["decision"]["id"]
    assert body["preview"] == "<PERSON>"

    status, _ = _request(
        "DELETE",
        f"{base}/review-sessions/{session_id}/decisions/user-added/{ua_id}",
        token=token,
    )
    assert status == 204
